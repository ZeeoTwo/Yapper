import sys 
sys.path.append('lib')
from pytube import YouTube
from moviepy.editor import VideoFileClip
import os
import assemblyai as aai
from api import API_KEY
from api import OPEN_AI_API_KEY
import openai
from docx import Document
from docx2pdf import convert
import tkinter as tk
from tkinter import ttk
from tkinter import messagebox
from tkinter import filedialog

openai.api_key = OPEN_AI_API_KEY
openai.Model.list()
aai.settings.api_key = API_KEY
class VideoNoteCreatorApp:
    def __init__(self, root):
        self.root = root
        self.uploaded_file_path = None
        self.root.title("YouTube Video Note Creator")

        self.upload_label = ttk.Label(root, text="Upload an audio or a video file:")
        self.upload_button = ttk.Button(root, text="upload", command=self.upload_file)        
        self.yt_link_label = ttk.Label(root, text="YouTube link:")
        self.yt_link_entry = ttk.Entry(root, width=50)

        self.lang_label = ttk.Label(root, text="Language of the video:")
        self.lang_entry_combobox = ttk.Combobox(root, values=["en", "en_au", "en_uk", "en_us", "es", "fr", "de", "it", "pt", "nl", "hi", "ja", "zh", "fi", "ko", "pl", "ru", "tr", "uk", "vi"])
        self.lang_entry_combobox.set("en")

        self.file_format_label = ttk.Label(root, text="Type of file for your note:")
        self.file_format_combobox = ttk.Combobox(root, values=["docx", "pdf", "txt", "md"])
        self.file_format_combobox.set("docx")

        self.note_directory_label = ttk.Label(root, text="Select note directory:")
        self.note_directory_button = ttk.Button(root, text="Select", command=self.select_note_directory)

        self.process_button = ttk.Button(root, text="Process Video", command=self.download_and_process)
        self.yt_link_entry.bind("<KeyRelease>", self.update_state)

        self.yt_link_label.grid(row=0, column=0, padx=10, pady=10)
        self.yt_link_entry.grid(row=0, column=1, columnspan=2, padx=10, pady=10)

        self.upload_label.grid(row=1,column=0,padx=10, pady=10)
        self.upload_button.grid(row=1, column=1, padx=10, pady=10)
        self.lang_label.grid(row=2, column=0, padx=10, pady=10)
        self.lang_entry_combobox.grid(row=2, column=1, padx=10, pady=10)

        self.file_format_label.grid(row=3, column=0, padx=10, pady=10)
        self.file_format_combobox.grid(row=3, column=1, padx=10, pady=10)

        self.note_directory_label.grid(row=4, column=0, padx=10, pady=10)
        self.note_directory_button.grid(row=4, column=1, padx=10, pady=10)

        self.process_button.grid(row=5, column=0, columnspan=3, padx=10, pady=10)
        
    def update_state(self,event):
        if self.yt_link_entry.get():
            self.upload_button.config(state="disabled")
        else:
            self.upload_button.config(state="normal")
    def reset_values(self):
        self.yt_link_entry.delete(0,tk.END)
        self.yt_link_entry.insert(0,"")
        self.lang_entry_combobox.set("en")
        self.file_format_combobox.set("docx")
        self.uploaded_file_path = ""

    def download_and_process(self):
        lang = self.lang_entry_combobox.get()
        file_format = self.file_format_combobox.get()
        if self.uploaded_file_path:
            self.yt_link_entry.config(state="disabled")
            self.upload_button.config(state="normal")
            transcript = self.transcribe_audio(self.uploaded_file_path,lang)
            if transcript: 
                note = self.get_summary_openai(transcript.text)
                if not self.create_note_file(file_format, note, self.uploaded_file_path):
                    messagebox.showinfo("Success", "Note file created successfully.")
                    self.yt_link_entry.config(state ="normal")
                    self.upload_button.config(state = "normal")
                    self.reset_values()

                else:
                    messagebox.showerror("Error", "Failed to create the note file.")

        if  self.yt_link_entry.get():
            self.yt_link_entry.config(state="normal")
            self.upload_button.config(state="disabled")
            yt_link = self.yt_link_entry.get()
            video_file, audio_url, file_name, lang, correct_file_format = self.download_video(yt_link, lang, file_format)
            if video_file and audio_url:
                self.extract_audio(video_file, audio_url)
                transcript = self.transcribe_audio(audio_url, lang)
                if transcript:
                    note = self.get_summary_openai(transcript.text)
                    self.create_note_file(correct_file_format, note, file_name)
                    if os.path.exists(audio_url):
                        os.remove(video_file)
                        os.remove(audio_url)
                        messagebox.showinfo("Success", "Note file created successfully.")
                        self.yt_link_entry.config(state ="normal")
                        self.upload_button.config(state = "normal")
                        self.reset_values()
                    else:
                        messagebox.showerror("Error", "Failed to create the note file.")

    def get_summary_openai(self,transcript):
        response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": "make a note from given transcript, state a title of the transcription, state the most important points of the transcription, add to that note a detailed explanation to the scientific topics or subjects included in transcription: " + transcript}])
        return response.choices[0].message.content

    def is_valid_youtube_url(self, url):
        try:
            yt = YouTube(url)
            return yt.title is not None
        except Exception:
            return False

    def download_video(self, yt_link, lang, file_format):
        while True:
            if not self.is_valid_youtube_url(yt_link):
                messagebox.showerror("Error", "Provided youtube link is either nonexisted or unavailable.")
                break
            else:
                youtube_obj = YouTube(yt_link)
                video_file = youtube_obj.video_id + ".mp4"
                audio_url = youtube_obj.video_id + ".mp3"
                stream = youtube_obj.streams.get_highest_resolution()
                stream.download(filename=video_file)
                return video_file, audio_url, youtube_obj.video_id, lang, file_format

    def create_note_file(self,file_format, note, file_name):
        if file_format == "docx":
            doc = Document()
            doc.add_paragraph(note)
            file_path = os.path.join(self.note_directory_button, file_name + ".docx")
            doc.save(file_path)
        if file_format == "pdf":
            doc = Document()
            doc.add_paragraph(note)
            file_path = os.path.join(self.note_directory_button, file_name + ".docx")
            doc.save(file_path)
            convert(file_path)
            if os.path.exists(file_path):
                os.remove(file_path)
        if file_format == "txt":
            file_path = os.path.join(self.note_directory_button, file_name + ".txt")
            with open(file_path, "w") as file_obj:
                file_obj.write(note)

    def extract_audio(self,video_file, audio_url):
        video = VideoFileClip(video_file)
        video.audio.write_audiofile(audio_url, codec="mp3")
        video.close()

    def transcribe_audio(self,audio_url, lang):
        cfg_transcription = aai.TranscriptionConfig(
            language_code=lang,
        )
        transcriber = aai.Transcriber(config=cfg_transcription)
        transcript = transcriber.transcribe(audio_url)
        if transcript.text == "none":
            messagebox.showerror("Error", "Failed to create the note file.")
            return None
        return transcript
    
    def upload_file(self):
        self.uploaded_file_path = filedialog.askopenfilename(
        title="Select a Video or Audio File",
        filetypes=[
            ("Video Files", "*.mts *.m2ts *.ts *.mov *.mp2 *.mp4 *.m4v *.mxf *.webm"),
            ("Audio Files", "*.8svx *.aac *.ac3 *.aif *.aiff *.alac *.amr *.ape *.au *.dss *.flac *.flv *.m4a *.m4b *.m4p *.m4r *.mp3 *.mpga *.ogg *.oga *.mogg *.opus *.qcp *.tta *.voc *.wav *.wma *.wv"),
        ])
        self.upload_button.config(state="disabled")
        self.yt_link_entry.config(state ="disabled")

    def select_note_directory(self):
        self.note_directory_button = filedialog.askdirectory(
        title="Select a directory Folder")

def main():
    app = tk.Tk()
    VideoNoteCreatorApp(app)
    app.mainloop()
if __name__ == "__main__":
    main()